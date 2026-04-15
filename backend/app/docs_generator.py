"""
Document generation utilities.
Generates User Guide (.docx) in Vietnamese and Application Introduction (.pdf).
"""

import os
from pathlib import Path

# ── Paths ────────────────────────────────────────────────────────────
_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
_USER_GUIDE_DIR = _PROJECT_ROOT / "User_Guide"
_DOCS_DIR = _PROJECT_ROOT / "docs"


def generate_user_guide_docx() -> str:
    """Generate User_Guide/User Guide.docx (Vietnamese) and return the file path."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _USER_GUIDE_DIR.mkdir(parents=True, exist_ok=True)
    filepath = _USER_GUIDE_DIR / "User Guide.docx"

    doc = Document()

    # ── Styles ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ── Title ────────────────────────────────────────────────────
    title = doc.add_heading(
        "Ứng dụng Từ điển Tiếng Anh sử dụng Radix Trie", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Báo cáo dự án – Cấu trúc dữ liệu và Giải thuật nâng cao"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 116, 139)

    # ── 1. Giới thiệu ───────────────────────────────────────────
    doc.add_heading("1. Giới thiệu", level=1)
    doc.add_paragraph(
        "Dự án xây dựng một ứng dụng từ điển Tiếng Anh sử dụng Radix Trie "
        "(cây Trie nén, còn gọi là Patricia Trie) làm cấu trúc dữ liệu lập "
        "chỉ mục chính. Ứng dụng cho phép người dùng thêm, xóa và tra cứu "
        "từ vựng thông qua giao diện web hiện đại, đồng thời trực quan hóa "
        "cấu trúc cây trong thời gian thực."
    )
    doc.add_paragraph(
        "Mục tiêu của dự án là chuyển hóa kiến thức lý thuyết về cấu trúc "
        "dữ liệu thành một hệ thống trực quan, có thể thao tác và có khả "
        "năng giải thích — giúp người học hiểu rõ cơ chế nén cạnh, tách nút "
        "và gộp nút thông qua tương tác thực tế."
    )

    # ── 2. Cơ sở lý thuyết ──────────────────────────────────────
    doc.add_heading("2. Cơ sở lý thuyết", level=1)

    doc.add_heading("2.1. Trie chuẩn", level=2)
    doc.add_paragraph(
        "Trie (phát âm \"try\") là cấu trúc cây dùng để lưu trữ chuỗi ký tự. "
        "Trong Trie chuẩn, mỗi nút đại diện cho một ký tự duy nhất; đường đi "
        "từ gốc đến một nút được đánh dấu kết thúc sẽ tạo thành một từ hoàn chỉnh."
    )
    doc.add_paragraph(
        "Ưu điểm: tra cứu nhanh O(m) với m là chiều dài từ. "
        "Nhược điểm: lãng phí bộ nhớ khi nhiều từ chia sẻ tiền tố dài, "
        "vì mỗi ký tự chiếm một nút riêng."
    )

    doc.add_heading("2.2. Radix Trie (Trie nén)", level=2)
    doc.add_paragraph(
        "Radix Trie cải tiến Trie chuẩn bằng cách gộp các chuỗi nút đơn con "
        "thành một cạnh duy nhất mang nhãn là chuỗi con (substring). Điều này "
        "giảm số lượng nút và cạnh, tiết kiệm bộ nhớ, đồng thời tăng tốc "
        "tra cứu vì mỗi bước so khớp nhiều ký tự cùng lúc."
    )
    doc.add_paragraph(
        "Ví dụ: ba từ \"cat\", \"car\", \"card\" trong Radix Trie được biểu diễn "
        "với cạnh \"ca\" gốc, sau đó phân nhánh thành \"t\", \"r\" và \"r→d\"."
    )

    # ── 3. Các thao tác ─────────────────────────────────────────
    doc.add_heading("3. Các thao tác trên Radix Trie", level=1)

    doc.add_heading("3.1. Thêm từ (Insert)", level=2)
    doc.add_paragraph(
        "Khi thêm từ mới, thuật toán duyệt cây từ gốc và so khớp với nhãn cạnh:"
    )
    doc.add_paragraph(
        "• Không có cạnh nào khớp: tạo cạnh mới mang toàn bộ phần còn lại.", style="List Bullet"
    )
    doc.add_paragraph(
        "• Nhãn cạnh khớp hoàn toàn: đi theo cạnh, tiếp tục với phần còn lại.", style="List Bullet"
    )
    doc.add_paragraph(
        "• Khớp một phần → Tách nút: cạnh hiện tại bị tách tại điểm khớp; "
        "phần chung trở thành cạnh mới dẫn đến nút trung gian; hai nhánh "
        "được tạo cho phần cũ và phần mới.", style="List Bullet"
    )

    doc.add_heading("3.2. Xóa từ (Delete)", level=2)
    doc.add_paragraph(
        "1. Tìm nút tương ứng bằng cách duyệt theo các cạnh.\n"
        "2. Bỏ đánh dấu nút kết thúc từ.\n"
        "3. Dọn dẹp cấu trúc cây:\n"
        "   – Nếu nút không còn con → xóa cạnh.\n"
        "   – Nếu nút cha chỉ còn một con và không phải kết thúc từ → gộp hai cạnh."
    )

    doc.add_heading("3.3. Tra cứu từ (Search)", level=2)
    doc.add_paragraph(
        "Tra cứu là thao tác chỉ đọc — không thay đổi cấu trúc cây. "
        "Thuật toán duyệt từ gốc, so khớp nhãn cạnh với từ cần tra, "
        "và trả về kết quả tìm thấy/không tìm thấy kèm đường đi đã duyệt."
    )

    # ── 4. Kiến trúc hệ thống ───────────────────────────────────
    doc.add_heading("4. Kiến trúc hệ thống", level=1)
    doc.add_paragraph(
        "Hệ thống gồm hai tầng chính giao tiếp qua HTTP REST API:"
    )

    # Table: Tech stack
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0]
    hdr.cells[0].text = "Thành phần"
    hdr.cells[1].text = "Công nghệ"
    data = [
        ("Frontend", "Next.js, TypeScript, Tailwind CSS, React Flow"),
        ("Backend", "Python 3.11+, FastAPI, Pydantic"),
        ("Cấu trúc dữ liệu", "Radix Trie (cài đặt thủ công)"),
        ("Lưu trữ", "Tệp JSON cục bộ (dictionary.json)"),
    ]
    for i, (col1, col2) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = col1
        row.cells[1].text = col2

    doc.add_paragraph()  # spacer

    # ── 5. Chức năng chính ───────────────────────────────────────
    doc.add_heading("5. Chức năng chính", level=1)

    # Table: Features
    feat_table = doc.add_table(rows=8, cols=3)
    feat_table.style = "Light Grid Accent 1"
    hdr = feat_table.rows[0]
    hdr.cells[0].text = "STT"
    hdr.cells[1].text = "Chức năng"
    hdr.cells[2].text = "Ý nghĩa"
    features = [
        ("1", "Thêm từ", "Chèn từ mới vào Radix Trie; tách nút nếu cần."),
        ("2", "Xóa từ", "Xóa từ khỏi cây; gộp nút nếu thừa."),
        ("3", "Tra cứu từ", "Duyệt cây và hiển thị đường đi."),
        ("4", "Trực quan hóa cây", "Hiển thị Radix Trie dạng cây tương tác."),
        ("5", "Lịch sử thao tác", "Ghi lại mọi thao tác kèm giải thích."),
        ("6", "Bảng từ vựng", "Hiển thị và lọc toàn bộ từ đã lưu."),
        ("7", "Load demo / Reset", "Nạp dữ liệu mẫu hoặc khôi phục rỗng."),
    ]
    for i, (stt, func, desc) in enumerate(features):
        row = feat_table.rows[i + 1]
        row.cells[0].text = stt
        row.cells[1].text = func
        row.cells[2].text = desc

    doc.add_paragraph()  # spacer

    # ── 6. Giao diện và quy trình sử dụng ───────────────────────
    doc.add_heading("6. Giao diện và quy trình sử dụng", level=1)
    doc.add_paragraph(
        "Giao diện chính được tổ chức thành các vùng chức năng rõ ràng: "
        "bảng điều khiển thao tác ở bên trái, vùng trực quan hóa Radix Trie "
        "bên phải, bảng từ vựng và lịch sử thao tác ở phía dưới."
    )
    doc.add_paragraph(
        "Quy trình thao tác chuẩn:"
    )
    doc.add_paragraph(
        "1. Khởi động backend và frontend (xem Mục 7).", style="List Number"
    )
    doc.add_paragraph(
        "2. Nhấn Load Demo để nạp dữ liệu mẫu, hoặc thêm từ thủ công.", style="List Number"
    )
    doc.add_paragraph(
        "3. Thực hiện thao tác thêm / xóa / tra cứu.", style="List Number"
    )
    doc.add_paragraph(
        "4. Quan sát cây cập nhật tức thì trên vùng trực quan hóa.", style="List Number"
    )
    doc.add_paragraph(
        "5. Kiểm tra kết quả chi tiết tại bảng kết quả thao tác.", style="List Number"
    )
    doc.add_paragraph(
        "6. Rê chuột vào dòng trong bảng từ vựng để làm nổi bật đường đi.", style="List Number"
    )

    # ── 7. Hướng dẫn cài đặt ────────────────────────────────────
    doc.add_heading("7. Hướng dẫn cài đặt và chạy", level=1)

    doc.add_heading("7.1. Yêu cầu", level=2)
    doc.add_paragraph("• Python 3.11 trở lên", style="List Bullet")
    doc.add_paragraph("• Node.js 18 trở lên", style="List Bullet")
    doc.add_paragraph("• npm (đi kèm Node.js)", style="List Bullet")

    doc.add_heading("7.2. Cài đặt Backend", level=2)
    doc.add_paragraph("Mở terminal tại thư mục gốc dự án:")
    p = doc.add_paragraph()
    run = p.add_run("cd backend\npip install -r requirements.txt")
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)

    doc.add_heading("7.3. Cài đặt Frontend", level=2)
    doc.add_paragraph("Mở terminal khác tại thư mục gốc dự án:")
    p = doc.add_paragraph()
    run = p.add_run("cd frontend\nnpm install")
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)

    doc.add_heading("7.4. Chạy ứng dụng", level=2)
    doc.add_paragraph("Bước 1 – Khởi động Backend:")
    p = doc.add_paragraph()
    run = p.add_run("cd backend\nuvicorn app.main:app --reload --port 8000")
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    doc.add_paragraph("API khả dụng tại http://localhost:8000")

    doc.add_paragraph("Bước 2 – Khởi động Frontend:")
    p = doc.add_paragraph()
    run = p.add_run("cd frontend\nnpm run dev")
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    doc.add_paragraph("Giao diện web khả dụng tại http://localhost:3000")

    # ── 8. Cấu trúc repository ──────────────────────────────────
    doc.add_heading("8. Cấu trúc repository", level=1)

    repo_table = doc.add_table(rows=8, cols=3)
    repo_table.style = "Light Grid Accent 1"
    hdr = repo_table.rows[0]
    hdr.cells[0].text = "Khu vực"
    hdr.cells[1].text = "Tệp / thư mục"
    hdr.cells[2].text = "Chức năng"
    repo_data = [
        ("Backend", "backend/app/radix_trie.py", "Cài đặt thuật toán Radix Trie"),
        ("Backend", "backend/app/routes.py", "Khai báo API endpoint"),
        ("Backend", "backend/app/storage.py", "Đọc/ghi trạng thái hệ thống"),
        ("Frontend", "frontend/src/app/page.tsx", "Quản lý state và luồng chính"),
        ("Frontend", "frontend/src/components/trie/", "Trực quan hóa Radix Trie"),
        ("Kiểm thử", "backend/tests/", "21 test cases cho Radix Trie"),
        ("Tài liệu", "User_Guide/, README.md", "Tài liệu thuyết minh"),
    ]
    for i, (area, path, desc) in enumerate(repo_data):
        row = repo_table.rows[i + 1]
        row.cells[0].text = area
        row.cells[1].text = path
        row.cells[2].text = desc

    doc.add_paragraph()

    # ── 9. Kiểm thử ─────────────────────────────────────────────
    doc.add_heading("9. Kiểm thử", level=1)
    doc.add_paragraph(
        "Dự án bao gồm 21 test cases cho cấu trúc Radix Trie, gồm:"
    )
    doc.add_paragraph("• 7 kiểm thử chèn (insert)", style="List Bullet")
    doc.add_paragraph("• 5 kiểm thử tra cứu (search)", style="List Bullet")
    doc.add_paragraph("• 6 kiểm thử xóa (delete)", style="List Bullet")
    doc.add_paragraph("• 3 kiểm thử trực quan hóa", style="List Bullet")

    doc.add_paragraph("Chạy kiểm thử:")
    p = doc.add_paragraph()
    run = p.add_run("cd backend\npython -m pytest tests/ -v")
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)

    # ── 10. Kết luận ─────────────────────────────────────────────
    doc.add_heading("10. Kết luận", level=1)
    doc.add_paragraph(
        "Dự án Ứng dụng Từ điển Tiếng Anh sử dụng Radix Trie minh họa "
        "thành công cách cấu trúc dữ liệu nâng cao có thể được ứng dụng "
        "trong một hệ thống thực tế. Việc kết hợp giữa cấu trúc dữ liệu "
        "hiệu quả (Radix Trie) với kiến trúc full-stack hiện đại (FastAPI + "
        "Next.js) tạo ra một công cụ trực quan, có giá trị học thuật — giúp "
        "người học không chỉ hiểu lý thuyết mà còn quan sát trực tiếp cơ chế "
        "tách nút, gộp nút và nén cạnh thông qua tương tác thực tế."
    )

    doc.save(str(filepath))
    return str(filepath)


def generate_app_intro_pdf() -> str:
    """Generate docs/Application_Introduction.pdf and return the file path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )

    _DOCS_DIR.mkdir(parents=True, exist_ok=True)
    filepath = _DOCS_DIR / "Application_Introduction.pdf"

    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=22,
        spaceAfter=12,
        textColor=HexColor("#1e293b"),
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=14,
        spaceBefore=18,
        spaceAfter=8,
        textColor=HexColor("#334155"),
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["BodyText"],
        fontSize=11,
        leading=16,
        spaceAfter=8,
    )

    elements: list = []

    # Title
    elements.append(Paragraph("English Dictionary Application", title_style))
    elements.append(
        Paragraph("Using Radix-Trie Data Structure", styles["Heading3"])
    )
    elements.append(Spacer(1, 20))

    # 1. Objective
    elements.append(Paragraph("1. Project Objective", heading_style))
    elements.append(
        Paragraph(
            "This project implements an English dictionary application that uses "
            "a Radix Trie (compressed Trie) as the primary indexing data structure. "
            "The goal is to demonstrate practical usage of advanced data structures "
            "in a full-stack web application, suitable for an academic Advanced "
            "Data Structures &amp; Algorithms course.",
            body_style,
        )
    )

    # 2. Key Features
    elements.append(Paragraph("2. Key Features", heading_style))
    features = [
        ["Feature", "Description"],
        ["Add Word", "Insert a new word with its meaning into the dictionary."],
        ["Delete Word", "Remove a word from the dictionary."],
        ["Search Word", "Look up a word and retrieve its meaning."],
        [
            "Trie Visualization",
            "Display the Radix Trie structure with interactive React Flow nodes.",
        ],
        [
            "Operation History",
            "Log all operations with their outcomes and structural effects.",
        ],
        [
            "Data Persistence",
            "Store entries in a local JSON file, loaded on startup.",
        ],
    ]
    t = Table(features, colWidths=[4 * cm, 12 * cm])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1e293b")),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 1), (-1, -1), 6),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cbd5e1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f8fafc"), HexColor("#ffffff")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.append(t)
    elements.append(Spacer(1, 12))

    # 3. Why Radix Trie
    elements.append(Paragraph("3. Why Radix Trie?", heading_style))
    elements.append(
        Paragraph(
            "A standard Trie stores one character per edge, which wastes memory "
            "when there are long chains of single-child nodes. A Radix Trie "
            "(also called a Patricia Trie or compressed Trie) compresses these "
            "chains by storing substrings on edges. This results in fewer nodes, "
            "lower memory usage, and faster lookups — making it ideal for "
            "dictionary applications where many words share common prefixes.",
            body_style,
        )
    )

    # 4. How It Works
    elements.append(Paragraph("4. How the Application Works", heading_style))
    elements.append(
        Paragraph(
            "<b>Backend:</b> A Python FastAPI server manages the Radix Trie in memory. "
            "Dictionary entries are persisted in a JSON file that is loaded on startup. "
            "The API exposes endpoints for adding, deleting, searching words, "
            "retrieving trie snapshots, and operation history.",
            body_style,
        )
    )
    elements.append(
        Paragraph(
            "<b>Frontend:</b> A Next.js (React) application provides a modern, "
            "responsive interface. Users can add, delete, or search words through "
            "intuitive panels. The Radix Trie is visualised interactively using "
            "React Flow with real-time updates after every operation.",
            body_style,
        )
    )

    # 5. Interface Summary
    elements.append(Paragraph("5. Interface Summary", heading_style))
    elements.append(
        Paragraph(
            "The web interface is divided into several panels: a header with live "
            "statistics, a tabbed operation panel (Add / Delete / Search), an "
            "interactive Radix Trie visualization powered by React Flow, "
            "an operation result panel with structural explanations, a filterable "
            "dictionary entries table, and an operation history log.",
            body_style,
        )
    )

    # 6. Tech Stack
    elements.append(Paragraph("6. Technology Stack", heading_style))
    tech = [
        ["Layer", "Technology"],
        ["Frontend", "Next.js, TypeScript, Tailwind CSS, React Flow (@xyflow/react)"],
        ["Backend", "Python 3.11+, FastAPI, Pydantic"],
        ["Data Structure", "Radix Trie (implemented from scratch)"],
        ["Persistence", "Local JSON file"],
    ]
    t2 = Table(tech, colWidths=[4 * cm, 12 * cm])
    t2.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1e293b")),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 1), (-1, -1), 6),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cbd5e1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f8fafc"), HexColor("#ffffff")]),
            ]
        )
    )
    elements.append(t2)
    elements.append(Spacer(1, 12))

    # 7. Conclusion
    elements.append(Paragraph("7. Conclusion", heading_style))
    elements.append(
        Paragraph(
            "This project successfully demonstrates how a Radix Trie can power "
            "a practical dictionary application. The combination of an efficient "
            "data structure with a modern full-stack architecture results in a "
            "performant, visually appealing, and educationally valuable tool "
            "that clearly illustrates the internal operations of compressed "
            "trie structures.",
            body_style,
        )
    )

    doc.build(elements)
    return str(filepath)


if __name__ == "__main__":
    print("Generating User Guide (.docx) — Vietnamese...")
    path1 = generate_user_guide_docx()
    print(f"  -> {path1}")

    print("Generating Application Introduction (.pdf)...")
    path2 = generate_app_intro_pdf()
    print(f"  -> {path2}")

    print("Done.")
